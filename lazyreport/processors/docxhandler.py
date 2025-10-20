from docx import Document
from docxcompose.composer import Composer
from typing import Dict, List
from pathlib import Path

class ReportService:
    def createVulnerabilityReport(self, templatePath: str, outputPath: str, replacements: Dict[str, str]) -> None:

        ## --- Replaces all occurrences of specified text patterns in a Word (.docx) file ---

        doc = Document(templatePath)

        ## The complex, formatting-preserving replacement logic is maintained here.
        ## Note: Due to its complexity and direct manipulation of the docx XML structure;
        ## Logic of this function remains largely unchanged to preserve functionality.

        from docx.shared import Pt

        def insertParagraphsAfter(doc, paragraph, texts, markerRun=None, prefixText=None, suffixText=None):
            parent = paragraph._element.getparent()
            idx = parent.index(paragraph._element)
            if texts:
                while paragraph.runs:
                    paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
                if prefixText:
                    prefixRun = paragraph.add_run(prefixText)
                    if markerRun:
                        prefixRun.bold, prefixRun.italic, prefixRun.underline, prefixRun.font.name, prefixRun.font.size, prefixRun.font.color.rgb = markerRun.bold, markerRun.italic, markerRun.underline, markerRun.font.name, markerRun.font.size, (markerRun.font.color.rgb if markerRun.font.color.rgb else None)
                run = paragraph.add_run(texts[0])
                if markerRun:
                    run.bold, run.italic, run.underline, run.font.name, run.font.size, run.font.color.rgb = markerRun.bold, markerRun.italic, markerRun.underline, markerRun.font.name, markerRun.font.size, (markerRun.font.color.rgb if markerRun.font.color.rgb else None)
                if suffixText:
                    suffixRun = paragraph.add_run(suffixText)
                    if markerRun:
                        suffixRun.bold, suffixRun.italic, suffixRun.underline, suffixRun.font.name, suffixRun.font.size, suffixRun.font.color.rgb = markerRun.bold, markerRun.italic, markerRun.underline, markerRun.font.name, markerRun.font.size, (markerRun.font.color.rgb if markerRun.font.color.rgb else None)
                paragraph.alignment, paragraph.style, paragraph.paragraph_format.space_after = paragraph.alignment, paragraph.style, Pt(6)
            
            for part in texts[1:]:
                newPara = doc.add_paragraph()
                newPara.alignment, newPara.style = paragraph.alignment, paragraph.style
                pfSrc, pfDst = paragraph.paragraph_format, newPara.paragraph_format
                pfDst.left_indent, pfDst.right_indent, pfDst.first_line_indent, pfDst.keep_together, pfDst.keep_with_next, pfDst.page_break_before, pfDst.widow_control, pfDst.space_before, pfDst.space_after, pfDst.line_spacing, pfDst.line_spacing_rule = pfSrc.left_indent, pfSrc.right_indent, pfSrc.first_line_indent, pfSrc.keep_together, pfSrc.keep_with_next, pfSrc.page_break_before, pfSrc.widow_control, pfSrc.space_before, pfSrc.space_after, pfSrc.line_spacing, pfSrc.line_spacing_rule
                run = newPara.add_run(part)
                if markerRun:
                    run.bold, run.italic, run.underline, run.font.name, run.font.size, run.font.color.rgb = markerRun.bold, markerRun.italic, markerRun.underline, markerRun.font.name, markerRun.font.size, (markerRun.font.color.rgb if markerRun.font.color.rgb else None)
                parent.insert(idx+1, newPara._element)
                idx += 1

        def performReplacement(container):
            for paragraph in container.paragraphs:
                for key, value in replacements.items():
                    fullText = ''.join(run.text for run in paragraph.runs)
                    if key not in fullText: continue
                    
                    start, end = fullText.find(key), fullText.find(key) + len(key)
                    charCount, runIndices = 0, []
                    for i, run in enumerate(paragraph.runs):
                        nextCount = charCount + len(run.text)
                        if charCount < end and nextCount > start: runIndices.append(i)
                        charCount = nextCount
                    
                    markerText = ''.join(paragraph.runs[i].text for i in runIndices)
                    markerRun = paragraph.runs[runIndices[0]] if runIndices else None

                    if '\n' in value:
                        prefixText = fullText[:start] if start > 0 else None
                        suffixText = fullText[end:] if end < len(fullText) else None
                        if runIndices:
                            paragraph.runs[runIndices[0]].text = markerText.replace(key, "")
                            for i in runIndices[1:]: paragraph.runs[i].text = ''
                        insertParagraphsAfter(doc, paragraph, [p.strip() for p in value.split('\n') if p.strip()], markerRun, prefixText, suffixText)
                    else:
                        replaced = markerText.replace(key, value)
                        if runIndices:
                            paragraph.runs[runIndices[0]].text = replaced
                            for i in runIndices[1:]: paragraph.runs[i].text = ''

        performReplacement(doc)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    performReplacement(cell)

        doc.save(outputPath)
        print(f"\n [✓] File saved in: {outputPath}")

    def mergeReports(self, filePaths: List[str], outputPath: str) -> None:
        ## --- Merges a list of .docx files into a single output file ---
        if len(filePaths) < 2:
            print(" Need at least 2 files to merge.")
            return

        master = Document(filePaths[0])
        composer = Composer(master)

        for filePath in filePaths[1:]:
            doc = Document(filePath)
            composer.append(doc)

        composer.save(outputPath)
        print(f" Merged {len(filePaths)} files into {outputPath}\n")
